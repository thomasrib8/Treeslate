import requests
from docx import Document
from tqdm import tqdm
import time
import os
import pandas as pd
import logging
import openai

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def create_glossary(api_key, name, source_lang, target_lang, glossary_path):
    """
    Crée un glossaire DeepL à partir d'un fichier CSV.

    :param api_key: Clé API DeepL.
    :param name: Nom du glossaire.
    :param source_lang: Langue source.
    :param target_lang: Langue cible.
    :param glossary_path: Chemin du fichier CSV contenant le glossaire.
    :return: ID du glossaire créé.
    """
    api_url = "https://api.deepl.com/v2/glossaries"
    with open(glossary_path, "r") as glossary_file:
        glossary_content = glossary_file.read()

    headers = {
        "Authorization": f"DeepL-Auth-Key {api_key}",
        "Content-Type": "application/x-www-form-urlencoded",
    }

    data = {
        "name": name,
        "source_lang": source_lang,
        "target_lang": target_lang,
        "entries_format": "csv",
        "entries": glossary_content,
    }

    response = requests.post(api_url, headers=headers, data=data)
    response_data = response.json()

    if response.status_code in (200, 201) and "glossary_id" in response_data:
        logger.info("Glossaire créé avec succès.")
        return response_data["glossary_id"]
    else:
        logger.error(f"Échec de la création du glossaire: {response.text}")
        raise Exception(f"Échec de la création du glossaire: {response.text}")

def translate_docx_with_deepl(api_key, input_file_path, output_file_path, target_language, source_language, glossary_id=None):
    """
    Traduit un fichier DOCX avec l'API DeepL.

    :param api_key: Clé API DeepL.
    :param input_file_path: Chemin du fichier d'entrée.
    :param output_file_path: Chemin du fichier traduit.
    :param target_language: Langue cible.
    :param source_language: Langue source.
    :param glossary_id: ID du glossaire (optionnel).
    """
    api_url = "https://api.deepl.com/v2/document"
    headers = {"Authorization": f"DeepL-Auth-Key {api_key}"}
    data = {"target_lang": target_language, "source_lang": source_language}

    if glossary_id:
        data["glossary_id"] = glossary_id

    with open(input_file_path, "rb") as file:
        upload_response = requests.post(api_url, headers=headers, data=data, files={"file": file})

    if upload_response.status_code != 200:
        raise Exception(f"Échec du téléchargement du document: {upload_response.text}")

    upload_data = upload_response.json()
    document_id = upload_data["document_id"]
    document_key = upload_data["document_key"]

    status_url = f"{api_url}/{document_id}"
    while True:
        status_response = requests.post(status_url, headers=headers, data={"document_key": document_key})
        status_data = status_response.json()
        if status_data["status"] == "done":
            break
        elif status_data["status"] == "error":
            raise Exception(f"Erreur de traduction: {status_data}")
        time.sleep(1)

    download_url = f"{api_url}/{document_id}/result"
    download_response = requests.post(download_url, headers=headers, data={"document_key": document_key})

    if download_response.status_code == 200:
        with open(output_file_path, "wb") as output_file:
            output_file.write(download_response.content)
    else:
        raise Exception(f"Échec du téléchargement du document traduit: {download_response.text}")

def improve_translation(input_file, glossary_path, output_file, language_level, source_language, target_language, group_size, model):
    """
    Améliore la traduction d'un fichier DOCX en utilisant ChatGPT.

    :param input_file: Fichier d'entrée à améliorer.
    :param glossary_path: Chemin vers le glossaire.
    :param output_file: Fichier de sortie amélioré.
    :param language_level: Niveau de langue souhaité.
    :param source_language: Langue source.
    :param target_language: Langue cible.
    :param group_size: Nombre de paragraphes à traiter ensemble.
    :param model: Modèle GPT à utiliser.
    """
    doc = Document(input_file)
    glossary = read_glossary(glossary_path)
    output_doc = Document()
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    logger.debug(f"Chargé {len(paragraphs)} paragraphes pour traitement.")
    with tqdm(total=len(paragraphs), desc="Traitement des paragraphes") as pbar:
        for i in range(0, len(paragraphs), group_size):
            group = paragraphs[i : i + group_size]
            improved_text = process_paragraphs(group, glossary, language_level, source_language, target_language, model)
            if improved_text:
                output_doc.add_paragraph(improved_text)
            else:
                logger.warning(f"Skipping group {i // group_size + 1} due to an error.")
            pbar.update(len(group))
    output_doc.save(output_file)
    logger.debug(f"Document amélioré enregistré dans {output_file}.")

def convert_excel_to_csv(excel_path, csv_path):
    """
    Convertit un fichier Excel en CSV.

    :param excel_path: Chemin du fichier Excel.
    :param csv_path: Chemin du fichier CSV résultant.
    :return: Chemin du fichier CSV.
    """
    df = pd.read_excel(excel_path, header=None)
    df.to_csv(csv_path, index=False, header=False)
    return csv_path

def read_glossary(glossary_path):
    """
    Lit un glossaire à partir d'un fichier Word (.docx).

    :param glossary_path: Chemin vers le fichier .docx contenant le glossaire.
    :return: Dictionnaire contenant les termes source et cible.
    """
    glossary = {}
    try:
        doc = Document(glossary_path)
        for paragraph in doc.paragraphs:
            if ":" in paragraph.text:
                source, target = paragraph.text.split(":", 1)
                glossary[source.strip()] = target.strip()
        logger.debug(f"Glossaire chargé avec succès depuis {glossary_path}.")
    except Exception as e:
        logger.error(f"Erreur lors de la lecture du glossaire : {e}")
        raise
    return glossary

def process_paragraphs(paragraphs, glossary, language_level, source_language, target_language, model):
    """
    Envoie les paragraphes à ChatGPT pour amélioration de la traduction.

    :param paragraphs: Liste de paragraphes à traiter.
    :param glossary: Glossaire à appliquer.
    :param language_level: Niveau de langue souhaité.
    :param source_language: Langue source.
    :param target_language: Langue cible.
    :param model: Modèle GPT à utiliser.
    :return: Texte amélioré.
    """
    logger.debug(f"Traitement des paragraphes avec le modèle {model}.")
    prompt = (
        f"Translate the following text from {source_language} to {target_language} "
        f"and improve its quality to match the '{language_level}' language level.\n"
        f"Use the glossary strictly when applicable: {glossary}.\n\n"
    )

    for para in paragraphs:
        prompt += f"{para}\n\n"

    try:
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a skilled translator and editor."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=2048,
            temperature=0.7,
        )
        return response["choices"][0]["message"]["content"].strip()
    except openai.error.RateLimitError as e:
        logger.error(f"Rate limit reached: {e}. Adding delay before retrying.")
        raise
    except Exception as e:
        logger.error(f"An error occurred with OpenAI API: {e}")
        raise
